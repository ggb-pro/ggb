"""Document parser: extract text from PDF, Word, Markdown, plain text."""

import os
from dataclasses import dataclass


@dataclass
class ParsedSection:
    content: str
    section_type: str = "text"  # text/table/code/heading
    level: int = 0  # heading level
    page_number: int | None = None


@dataclass
class ParsedDocument:
    title: str
    sections: list[ParsedSection]
    raw_text: str
    page_count: int | None = None
    language: str = "zh"


def parse_file(file_path: str, mime_type: str | None = None) -> ParsedDocument:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_path)
    elif ext in (".md", ".markdown"):
        return _parse_markdown(file_path)
    elif ext in (".txt", ".text"):
        return _parse_text(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def _parse_pdf(file_path: str) -> ParsedDocument:
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    sections = []
    full_text = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        if text.strip():
            sections.append(ParsedSection(content=text.strip(), page_number=page_num + 1))
            full_text.append(text)

    doc.close()
    title = os.path.basename(file_path)
    if sections:
        first_line = sections[0].content.split("\n")[0][:100]
        title = first_line if first_line else title

    return ParsedDocument(
        title=title,
        sections=sections,
        raw_text="\n".join(full_text),
        page_count=len(doc),
    )


def _parse_docx(file_path: str) -> ParsedDocument:
    from docx import Document

    doc = Document(file_path)
    sections = []
    full_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Detect heading level from style
        style_name = para.style.name if para.style else ""
        level = 0
        section_type = "text"
        if "Heading" in style_name:
            try:
                level = int(style_name.replace("Heading ", "").replace("Heading", ""))
                section_type = "heading"
            except ValueError:
                pass
        sections.append(ParsedSection(content=text, section_type=section_type, level=level))
        full_text.append(text)

    title = os.path.basename(file_path)
    if sections and sections[0].section_type == "heading":
        title = sections[0].content[:100]

    return ParsedDocument(title=title, sections=sections, raw_text="\n".join(full_text))


def _parse_markdown(file_path: str) -> ParsedDocument:
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    sections = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        level = 0
        section_type = "text"
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            section_type = "heading"
        sections.append(ParsedSection(content=stripped, section_type=section_type, level=level))

    title = os.path.basename(file_path)
    if sections and sections[0].section_type == "heading":
        title = sections[0].content.lstrip("# ").strip()

    return ParsedDocument(title=title, sections=sections, raw_text=text)


def _parse_text(file_path: str) -> ParsedDocument:
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    sections = [ParsedSection(content=p.strip()) for p in text.split("\n\n") if p.strip()]
    return ParsedDocument(title=os.path.basename(file_path), sections=sections, raw_text=text)
